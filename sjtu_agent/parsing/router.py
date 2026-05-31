from __future__ import annotations

import csv
import html
import importlib.util
import inspect
import json
import re
import tempfile
import zipfile
from html.parser import HTMLParser
from pathlib import Path
from typing import Callable
from xml.etree import ElementTree as ET


TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".ini", ".cfg", ".conf",
    ".yaml", ".yml", ".toml", ".json", ".jsonl", ".xml", ".csv", ".tsv",
    ".py", ".java", ".c", ".cpp", ".h", ".hpp", ".js", ".ts", ".tsx", ".jsx",
    ".sql", ".tex",
}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
AUDIO_SUFFIXES = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}
VIDEO_SUFFIXES = {".mp4", ".mov", ".mkv", ".avi", ".webm"}


def _has_module(name: str) -> bool:
    return importlib.util.find_spec(name) is not None


def _truncate(text: str, max_chars: int) -> tuple[str, bool]:
    if max_chars <= 0:
        return "", bool(text)
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _result_ok(
    path: Path,
    parser: str,
    content: str,
    max_chars: int,
    warnings: list[str] | None = None,
    meta: dict | None = None,
) -> dict:
    body, truncated = _truncate(content, max_chars)
    return {
        "ok": True,
        "file": path.name,
        "path": str(path.resolve()),
        "suffix": path.suffix.lower(),
        "parser": parser,
        "truncated": truncated,
        "content": body,
        "warnings": warnings or [],
        "meta": meta or {},
    }


def _result_err(path: Path, error: str, warnings: list[str] | None = None) -> dict:
    return {
        "ok": False,
        "file": path.name,
        "path": str(path.resolve()),
        "suffix": path.suffix.lower(),
        "error": error,
        "warnings": warnings or [],
    }


def _append_warning(result: dict, warning: str) -> dict:
    warnings = result.get("warnings")
    if not isinstance(warnings, list):
        warnings = []
    if warning not in warnings:
        warnings.append(warning)
    result["warnings"] = warnings
    return result


class _StripHtml(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self.parts.append(data)


def _parse_pdf_pypdf(path: Path, max_chars: int, start_page: int = 1) -> dict:
    import pypdf

    reader = pypdf.PdfReader(str(path))
    total_pages = len(reader.pages)
    if total_pages == 0:
        return _result_ok(path, "pypdf", "", max_chars, meta={"total_pages": 0, "pages_read": ""})

    s = max(1, int(start_page))
    lines: list[str] = []
    chars = 0
    end_page = s - 1
    for page_idx in range(s - 1, total_pages):
        page_num = page_idx + 1
        text = (reader.pages[page_idx].extract_text() or "").strip()
        if not text:
            continue
        block = f"【第 {page_num} 页】\n{text}\n\n"
        if chars + len(block) > max_chars:
            remain = max(0, max_chars - chars)
            if remain > 0:
                lines.append(block[:remain])
                chars += remain
            end_page = page_num
            break
        lines.append(block)
        chars += len(block)
        end_page = page_num

    return {
        "ok": True,
        "file": path.name,
        "path": str(path.resolve()),
        "suffix": ".pdf",
        "parser": "pypdf",
        "truncated": chars >= max_chars,
        "content": "".join(lines).strip(),
        "warnings": [],
        "meta": {
            "total_pages": total_pages,
            "pages_read": f"{s}-{end_page}" if end_page >= s else "",
        },
    }


def _parse_html(path: Path, max_chars: int) -> dict:
    raw = path.read_text(encoding="utf-8", errors="replace")
    parser = _StripHtml()
    parser.feed(raw)
    text = html.unescape(" ".join(parser.parts))
    text = re.sub(r"\s{3,}", "\n\n", text).strip()
    return _result_ok(path, "html_parser", text, max_chars)


def _parse_text(path: Path, max_chars: int) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".json":
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            text = json.dumps(data, ensure_ascii=False, indent=2)
            return _result_ok(path, "json", text, max_chars)
        except Exception:
            pass

    if suffix in {".csv", ".tsv"}:
        delim = "," if suffix == ".csv" else "\t"
        lines: list[str] = []
        with path.open("r", encoding="utf-8", errors="replace", newline="") as fp:
            reader = csv.reader(fp, delimiter=delim)
            for i, row in enumerate(reader):
                lines.append(" | ".join(c.strip() for c in row))
                if i >= 500:
                    lines.append("... (rows truncated)")
                    break
        return _result_ok(path, "csv", "\n".join(lines), max_chars)

    text = path.read_text(encoding="utf-8", errors="replace")
    return _result_ok(path, "text", text, max_chars)


def _parse_docx(path: Path, max_chars: int) -> dict:
    from docx import Document

    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(lines).strip()
    return _result_ok(path, "python-docx", text, max_chars)


def _parse_markitdown(path: Path, max_chars: int) -> dict:
    from markitdown import MarkItDown  # type: ignore

    md = MarkItDown()
    out = md.convert(str(path))
    content = getattr(out, "text_content", None)
    if not isinstance(content, str):
        content = str(out)
    return _result_ok(path, "markitdown", content.strip(), max_chars)


def _parse_zip(path: Path, max_chars: int) -> dict:
    entries: list[str] = []
    with zipfile.ZipFile(path, "r") as zf:
        names = zf.namelist()
        for idx, name in enumerate(names):
            entries.append(name)
            if idx >= 199:
                entries.append("... (entries truncated)")
                break
    content = "ZIP archive entries:\n" + "\n".join(entries)
    return _result_ok(path, "zip_index", content, max_chars, meta={"entry_count": len(entries)})


def _parse_pptx_xml(path: Path, max_chars: int) -> dict:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    slides: list[tuple[int, str]] = []
    with zipfile.ZipFile(path, "r") as zf:
        slide_names = sorted(
            name
            for name in zf.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        for name in slide_names:
            m = re.search(r"slide(\d+)\.xml$", name)
            slide_idx = int(m.group(1)) if m else len(slides) + 1
            raw = zf.read(name)
            root = ET.fromstring(raw)
            texts = [t.text.strip() for t in root.findall(".//a:t", ns) if t.text and t.text.strip()]
            if texts:
                slides.append((slide_idx, "\n".join(texts)))

    if not slides:
        return _result_ok(path, "pptx_xml", "", max_chars, warnings=["No slide text found"])

    content = "\n\n".join(f"【第 {idx} 页】\n{text}" for idx, text in slides)
    return _result_ok(path, "pptx_xml", content, max_chars, meta={"slide_count": len(slides)})


def _make_paddle_ocr():
    from paddleocr import PaddleOCR  # type: ignore

    params = inspect.signature(PaddleOCR.__init__).parameters
    kwargs: dict = {}
    if "lang" in params:
        kwargs["lang"] = "ch"
    # PaddleOCR 2.x
    if "use_angle_cls" in params:
        kwargs["use_angle_cls"] = True
        if "show_log" in params:
            kwargs["show_log"] = False
    # PaddleOCR 3.x
    else:
        if "use_textline_orientation" in params:
            kwargs["use_textline_orientation"] = True
        if "use_doc_orientation_classify" in params:
            kwargs["use_doc_orientation_classify"] = False
        if "use_doc_unwarping" in params:
            kwargs["use_doc_unwarping"] = False
    return PaddleOCR(**kwargs)


def _extract_texts_from_ocr_result(obj) -> list[str]:
    texts: list[str] = []

    def _walk(x):
        if x is None:
            return
        if isinstance(x, str):
            t = x.strip()
            if t:
                texts.append(t)
            return
        if isinstance(x, dict):
            # PaddleOCR 3.x often exposes rec_texts/rec_text under res.
            for key in ("rec_texts", "rec_text", "text"):
                val = x.get(key)
                if isinstance(val, str):
                    t = val.strip()
                    if t:
                        texts.append(t)
                elif isinstance(val, (list, tuple)):
                    for item in val:
                        _walk(item)
            for v in x.values():
                _walk(v)
            return
        if isinstance(x, (list, tuple)):
            # PaddleOCR 2.x line format: [poly, (text, score)]
            if len(x) >= 2 and isinstance(x[1], (list, tuple)) and x[1]:
                maybe_txt = x[1][0]
                if isinstance(maybe_txt, str):
                    t = maybe_txt.strip()
                    if t:
                        texts.append(t)
            for item in x:
                _walk(item)
            return
        for attr in ("res", "rec_texts", "rec_text"):
            if hasattr(x, attr):
                try:
                    _walk(getattr(x, attr))
                except Exception:
                    pass

    _walk(obj)
    seen: set[str] = set()
    dedup: list[str] = []
    for t in texts:
        if t in seen:
            continue
        seen.add(t)
        dedup.append(t)
    return dedup


def _run_paddle_ocr_texts(ocr, image_path: str) -> list[str]:
    try:
        result = ocr.ocr(image_path, cls=True)
    except TypeError:
        result = ocr.ocr(image_path)
    texts = _extract_texts_from_ocr_result(result)
    if texts:
        return texts
    if hasattr(ocr, "predict"):
        result2 = ocr.predict(image_path)
        return _extract_texts_from_ocr_result(result2)
    return []


def _parse_pdf_ocr(path: Path, max_chars: int, start_page: int = 1) -> dict:
    import pypdfium2 as pdfium  # type: ignore

    with pdfium.PdfDocument(str(path)) as doc:
        total_pages = len(doc)
        if total_pages == 0:
            return _result_ok(path, "pdf_ocr", "", max_chars, warnings=["PDF has zero pages"], meta={"total_pages": 0})

        ocr = _make_paddle_ocr()
        blocks: list[str] = []
        start_idx = max(0, int(start_page) - 1)
        for i in range(start_idx, total_pages):
            page_no = i + 1
            page = doc[i]
            bmp = None
            pil = None
            try:
                bmp = page.render(scale=2.0)
                pil = bmp.to_pil()
                with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as fp:
                    pil.save(fp.name)
                    page_texts = _run_paddle_ocr_texts(ocr, fp.name)
            finally:
                try:
                    if pil is not None:
                        pil.close()
                except Exception:
                    pass
                try:
                    if bmp is not None:
                        bmp.close()
                except Exception:
                    pass
                try:
                    page.close()
                except Exception:
                    pass

            if page_texts:
                blocks.append(f"[Page {page_no}]\n" + "\n".join(page_texts))

        return _result_ok(
            path,
            "pdf_ocr",
            "\n\n".join(blocks).strip(),
            max_chars,
            meta={"total_pages": total_pages, "pages_read": f"{start_idx + 1}-{total_pages}"},
        )


def _parse_pptx_ocr_images(path: Path, max_chars: int) -> dict:
    media_names: list[str] = []
    with zipfile.ZipFile(path, "r") as zf:
        for name in zf.namelist():
            low = name.lower()
            if low.startswith("ppt/media/") and low.endswith((".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff")):
                media_names.append(name)

        if not media_names:
            return _result_ok(path, "pptx_ocr", "", max_chars, warnings=["No PPT media images found"])

        ocr = _make_paddle_ocr()
        blocks: list[str] = []
        for idx, name in enumerate(media_names, start=1):
            raw = zf.read(name)
            suffix = Path(name).suffix or ".png"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as fp:
                fp.write(raw)
                fp.flush()
                texts = _run_paddle_ocr_texts(ocr, fp.name)
            if texts:
                blocks.append(f"[Media {idx}: {Path(name).name}]\n" + "\n".join(texts))

    return _result_ok(path, "pptx_ocr", "\n\n".join(blocks).strip(), max_chars, meta={"media_count": len(media_names)})


def _parse_image_paddleocr(path: Path, max_chars: int) -> dict:
    ocr = _make_paddle_ocr()
    texts = _run_paddle_ocr_texts(ocr, str(path))
    return _result_ok(path, "paddleocr", "\n".join(texts), max_chars)


def _parse_image_stub(path: Path, max_chars: int, reason: str) -> dict:
    return _result_ok(
        path,
        "image_stub",
        f"[image] {path.name}\n{reason}",
        max_chars,
        warnings=[reason],
        meta={"media_type": "image"},
    )


def _parse_audio_whisper(path: Path, max_chars: int) -> dict:
    import whisper  # type: ignore

    model = whisper.load_model("base")
    out = model.transcribe(str(path))
    text = str(out.get("text", "")).strip()
    return _result_ok(path, "whisper", text, max_chars)


def _parse_audio_stub(path: Path, max_chars: int, reason: str) -> dict:
    return _result_ok(
        path,
        "audio_stub",
        f"[audio] {path.name}\n{reason}",
        max_chars,
        warnings=[reason],
        meta={"media_type": "audio"},
    )


def _parse_video_stub(path: Path, max_chars: int) -> dict:
    reason = "Video parsing is not enabled. Extract audio first, then parse the audio file."
    return _result_ok(
        path,
        "video_stub",
        f"[video] {path.name}\n{reason}",
        max_chars,
        warnings=[reason],
        meta={"media_type": "video"},
    )


def _maybe_parse_docling(path: Path, max_chars: int) -> dict:
    if not _has_module("docling"):
        return _result_err(path, "docling backend is not installed")
    reason = "docling backend is detected but not wired in this runtime build."
    return _result_err(path, reason, warnings=[reason])


def _maybe_parse_mineru(path: Path, max_chars: int) -> dict:
    if not _has_module("mineru"):
        return _result_err(path, "mineru backend is not installed")
    reason = "mineru backend is detected but not wired in this runtime build."
    return _result_err(path, reason, warnings=[reason])


def parse_file(
    file_path: str,
    max_chars: int = 8000,
    start_page: int = 1,
    strategy: str = "auto",
) -> dict:
    path = Path(file_path)
    if not path.exists():
        return {"ok": False, "error": f"文件不存在: {file_path}"}
    if not path.is_file():
        return {"ok": False, "error": f"路径不是文件: {file_path}"}

    suffix = path.suffix.lower()
    selected = (strategy or "auto").strip().lower()
    max_chars = int(max_chars) if int(max_chars) > 0 else 8000

    def _run(fn: Callable[[], dict]) -> dict:
        try:
            return fn()
        except Exception as e:
            return _result_err(path, str(e))

    # explicit strategy
    if selected == "legacy":
        return _result_err(path, "legacy strategy should be handled by caller")
    if selected == "markitdown":
        if not _has_module("markitdown"):
            return _result_err(path, "markitdown backend is not installed")
        return _run(lambda: _parse_markitdown(path, max_chars))
    if selected == "docling":
        return _run(lambda: _maybe_parse_docling(path, max_chars))
    if selected == "mineru":
        return _run(lambda: _maybe_parse_mineru(path, max_chars))
    if selected == "paddleocr":
        if not _has_module("paddleocr"):
            return _result_err(path, "paddleocr backend is not installed")
        return _run(lambda: _parse_image_paddleocr(path, max_chars))
    if selected == "whisper":
        if not _has_module("whisper"):
            return _result_err(path, "whisper backend is not installed")
        return _run(lambda: _parse_audio_whisper(path, max_chars))
    if selected == "pdf_ocr":
        if not _has_module("paddleocr") or not _has_module("pypdfium2"):
            return _result_err(path, "pdf_ocr backend is not installed (requires paddleocr + pypdfium2)")
        return _run(lambda: _parse_pdf_ocr(path, max_chars, start_page=start_page))

    # auto routing
    if suffix == ".pdf":
        base = _run(lambda: _parse_pdf_pypdf(path, max_chars, start_page=start_page))
        if not base.get("ok"):
            return base
        if str(base.get("content", "")).strip():
            return base
        if _has_module("paddleocr") and _has_module("pypdfium2"):
            ocr = _run(lambda: _parse_pdf_ocr(path, max_chars, start_page=start_page))
            if ocr.get("ok") and str(ocr.get("content", "")).strip():
                return _append_warning(ocr, "text_extraction_empty_fallback_to_pdf_ocr")
            return _append_warning(base, "text_extraction_empty_and_pdf_ocr_no_text")
        return _append_warning(base, "PDF OCR backend missing: install paddleocr and pypdfium2, then retry.")
    if suffix in {".html", ".htm"}:
        return _run(lambda: _parse_html(path, max_chars))
    if suffix in {".docx", ".doc"}:
        if _has_module("docx"):
            return _run(lambda: _parse_docx(path, max_chars))
    if suffix in TEXT_SUFFIXES:
        return _run(lambda: _parse_text(path, max_chars))
    if suffix == ".zip":
        return _run(lambda: _parse_zip(path, max_chars))
    if suffix == ".pptx":
        base = _run(lambda: _parse_pptx_xml(path, max_chars))
        if not base.get("ok"):
            return base
        if str(base.get("content", "")).strip():
            return base
        if _has_module("paddleocr"):
            ocr = _run(lambda: _parse_pptx_ocr_images(path, max_chars))
            if ocr.get("ok") and str(ocr.get("content", "")).strip():
                return _append_warning(ocr, "slide_text_empty_fallback_to_pptx_ocr")
            return _append_warning(base, "slide_text_empty_and_pptx_ocr_no_text")
        return _append_warning(base, "PPT OCR backend missing: install paddleocr, then retry.")
    if suffix in IMAGE_SUFFIXES:
        if _has_module("paddleocr"):
            return _run(lambda: _parse_image_paddleocr(path, max_chars))
        return _parse_image_stub(path, max_chars, "OCR backend missing: install paddleocr to extract text.")
    if suffix in AUDIO_SUFFIXES:
        if _has_module("whisper"):
            return _run(lambda: _parse_audio_whisper(path, max_chars))
        return _parse_audio_stub(path, max_chars, "ASR backend missing: install whisper for transcription.")
    if suffix in VIDEO_SUFFIXES:
        return _parse_video_stub(path, max_chars)

    # optional catch-all
    if _has_module("markitdown"):
        return _run(lambda: _parse_markitdown(path, max_chars))
    return _result_err(path, f"不支持的文件类型: {suffix or '[no extension]'}")


def parse_files(
    file_paths: list[str],
    per_file_max_chars: int = 4000,
    total_max_chars: int = 12000,
    start_page: int = 1,
    strategy: str = "auto",
) -> dict:
    if not isinstance(file_paths, list) or not file_paths:
        return {"ok": False, "error": "file_paths 不能为空"}

    merged: list[str] = []
    items: list[dict] = []
    failures: list[dict] = []
    total_chars = 0

    for raw in file_paths:
        r = parse_file(
            file_path=str(raw),
            max_chars=per_file_max_chars,
            start_page=start_page,
            strategy=strategy,
        )
        if not r.get("ok"):
            failures.append({"file_path": str(raw), "error": r.get("error", "parse failed")})
            items.append(r)
            continue

        content = str(r.get("content", ""))
        if not content:
            items.append(r)
            continue

        block = f"===== {r.get('file', raw)} =====\n{content}\n"
        if total_chars + len(block) > total_max_chars:
            remain = max(0, total_max_chars - total_chars)
            if remain > 0:
                merged.append(block[:remain])
                total_chars += remain
            items.append({**r, "truncated": True})
            break

        merged.append(block)
        total_chars += len(block)
        items.append(r)

    return {
        "ok": True,
        "strategy": strategy,
        "count": len(file_paths),
        "success_count": len([i for i in items if i.get("ok")]),
        "failure_count": len(failures),
        "failures": failures,
        "truncated": total_chars >= total_max_chars,
        "content": "\n".join(merged).strip(),
        "items": items,
    }
